#!/usr/bin/env python3
"""
Product Document Generator
Reads product data from products_data.json and generates individual Word documents for each product.
"""

import json
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# Load the product data from JSON file
with open('products_data.json', 'r') as f:
    products = json.load(f)

def clean_filename(name):
    """Clean product name for use as filename"""
    # Remove special characters and replace spaces with underscores
    cleaned = re.sub(r'[^\w\s-]', '', name)
    cleaned = re.sub(r'[-\s]+', '_', cleaned)
    return cleaned.strip('_')

def add_formatted_text(paragraph, text, bold=False, size=None):
    """Add formatted text to a paragraph"""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return run

def is_title_field(key):
    """Check if a field should be formatted as a title (bold and larger)"""
    title_keywords = ['title', 'name', 'sku']
    return any(keyword in key.lower() for keyword in title_keywords)

def format_field_name(key):
    """Format field names for display (remove underscores, capitalize)"""
    # Skip certain technical fields
    skip_fields = ['Main category', 'Cost', 'Landing Cost', 'Bag size', 'Bag Price', 'price/lb']
    if key in skip_fields:
        return None
    
    # Clean up field names
    formatted = key.replace('_', ' ').replace('/', ' / ')
    # Capitalize each word
    formatted = ' '.join(word.capitalize() for word in formatted.split())
    return formatted

def add_table(doc, title, data_dict):
    """Add a formatted table to the document"""
    # Add table title
    heading = doc.add_heading(level=2)
    heading_run = heading.runs[0] if heading.runs else heading.add_run()
    heading_run.text = title
    heading_run.font.size = Pt(14)
    heading_run.bold = True
    
    # Create table
    table = doc.add_table(rows=len(data_dict), cols=2)
    table.style = 'Table Grid'
    
    # Populate table
    for i, (key, value) in enumerate(data_dict.items()):
        row = table.rows[i]
        
        # Format key cell (left column)
        key_cell = row.cells[0]
        key_paragraph = key_cell.paragraphs[0]
        key_run = key_paragraph.add_run(format_field_name(key) or key)
        key_run.bold = True
        key_run.font.size = Pt(11)
        
        # Format value cell (right column)
        value_cell = row.cells[1]
        value_paragraph = value_cell.paragraphs[0]
        value_run = value_paragraph.add_run(str(value))
        value_run.font.size = Pt(11)
    
    # Add spacing after table
    doc.add_paragraph()

def create_product_document(product_data):
    """Create a Word document for a single product"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Product Title (Main heading)
    title_paragraph = doc.add_heading(level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0] if title_paragraph.runs else title_paragraph.add_run()
    title_run.text = product_data.get('Title', 'Unknown Product')
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    # Add some spacing
    doc.add_paragraph()
    
    # Define table 1 fields (Growing Conditions)
    table1_fields = [
        "Sun Requirements (Full Sun, Full Sun to Partial Shade, Shade)",
        "Soil Preference", 
        "Soil pH",
        "Days to Maturity",
        "Height when mature",
        "Seeding rate",
        "Planting Depth"
    ]
    
    # Define table 2 fields (Plant Characteristics) 
    table2_fields = [
        "Sun/Shade",
        "Height when Mature", 
        "Seeding Rate",
        "Uses",
        "Color",
        "Water",
        "Native/Introduced",
        "Life Form"
    ]
    
    # Track which fields we've processed
    processed_fields = set(['Title'])
    
    # Process fields in order, creating tables and other content
    for key, value in product_data.items():
        if not value or value == '' or key in processed_fields:
            continue
            
        formatted_key = format_field_name(key)
        if formatted_key is None:
            processed_fields.add(key)
            continue
        
        # Check if we should create Table 1 (Growing Conditions)
        if key in table1_fields and key not in processed_fields:
            table1_data = {}
            for field in table1_fields:
                if field in product_data and product_data[field]:
                    table1_data[field] = product_data[field]
                    processed_fields.add(field)
            
            if table1_data:
                add_table(doc, "Growing Conditions", table1_data)
            continue
        
        # Handle SKU and Scientific Name
        if key in ['SKU', 'Scientific Name / mix %']:
            if is_title_field(key):
                heading = doc.add_heading(level=2)
                heading_run = heading.runs[0] if heading.runs else heading.add_run()
                heading_run.text = str(value)
                heading_run.font.size = Pt(16)
            else:
                p = doc.add_paragraph()
                label_run = p.add_run(f"{formatted_key}: ")
                label_run.bold = True
                label_run.font.size = Pt(12)
                value_run = p.add_run(str(value))
                value_run.font.size = Pt(11)
            processed_fields.add(key)
            continue
        
        # Handle description
        if 'description' in key.lower() or 'what is' in key.lower():
            p = doc.add_paragraph()
            label_run = p.add_run(f"{formatted_key}: ")
            label_run.bold = True
            label_run.font.size = Pt(12)
            value_run = p.add_run(str(value))
            value_run.font.size = Pt(12)
            p.space_after = Pt(12)
            processed_fields.add(key)
            continue
        
        # Handle "Why chose this product" sections
        if 'why chose this product title' in key.lower():
            # Find the corresponding content
            title_num = key.split()[-1]  # Get the number
            content_key = f"Why chose this product {title_num}"
            
            if content_key in product_data and product_data[content_key]:
                heading = doc.add_heading(level=3)
                heading_run = heading.runs[0] if heading.runs else heading.add_run()
                heading_run.text = str(value)
                heading_run.font.size = Pt(14)
                
                p = doc.add_paragraph()
                content_run = p.add_run(str(product_data[content_key]))
                content_run.font.size = Pt(11)
                p.space_after = Pt(6)
                
                processed_fields.add(key)
                processed_fields.add(content_key)
            continue
        
        # Check if we should create Table 2 (Plant Characteristics) after "Why chose" sections
        if key in table2_fields and key not in processed_fields:
            # Check if we've processed all "Why chose" sections first
            why_chose_done = True
            for i in range(1, 6):
                title_key = f"Why chose this product Title {i}"
                if title_key in product_data and title_key not in processed_fields:
                    why_chose_done = False
                    break
            
            if why_chose_done:
                table2_data = {}
                for field in table2_fields:
                    if field in product_data and product_data[field]:
                        table2_data[field] = product_data[field]
                        processed_fields.add(field)
                
                if table2_data:
                    add_table(doc, "Plant Characteristics", table2_data)
                continue
        
        # Handle Planting Guide sections
        if 'planting guide step' in key.lower():
            step_num = key.split()[-1]
            heading = doc.add_heading(level=3)
            heading_run = heading.runs[0] if heading.runs else heading.add_run()
            heading_run.text = f"Step {step_num}"
            heading_run.font.size = Pt(12)
            
            p = doc.add_paragraph()
            content_run = p.add_run(str(value))
            content_run.font.size = Pt(11)
            p.space_after = Pt(6)
            processed_fields.add(key)
            continue
        
        # Handle FAQ sections
        if key.startswith('FAQ '):
            faq_num = key.split()[-1]
            p = doc.add_paragraph()
            question_run = p.add_run(f"Q{faq_num}: ")
            question_run.bold = True
            question_run.font.size = Pt(11)
            answer_run = p.add_run(str(value))
            answer_run.font.size = Pt(11)
            p.space_after = Pt(6)
            processed_fields.add(key)
            continue
        
        # Skip fields that are part of "Why chose this product" content (already handled above)
        if 'why chose this product' in key.lower() and not 'title' in key.lower():
            processed_fields.add(key)
            continue
        
        # Handle any remaining fields
        if key not in processed_fields:
            if is_title_field(key):
                heading = doc.add_heading(level=2)
                heading_run = heading.runs[0] if heading.runs else heading.add_run()
                heading_run.text = str(value)
                heading_run.font.size = Pt(16)
            else:
                p = doc.add_paragraph()
                label_run = p.add_run(f"{formatted_key}: ")
                label_run.bold = True
                label_run.font.size = Pt(12)
                value_run = p.add_run(str(value))
                value_run.font.size = Pt(11)
            processed_fields.add(key)
    
    return doc

def main():
    """Main function to generate all product documents"""
    print(f"Found {len(products)} products to process...")
    
    # Create output directory
    import os
    output_dir = "Product_Documents"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Generate document for each product
    for i, product in enumerate(products, 1):
        product_name = product.get('Title', f'Product_{i}')
        filename = clean_filename(product_name)
        
        print(f"Generating document for: {product_name}")
        
        # Create the document
        doc = create_product_document(product)
        
        # Save the document
        output_path = os.path.join(output_dir, f"{filename}.docx")
        doc.save(output_path)
        
        print(f"  Saved: {output_path}")
    
    print(f"\nCompleted! Generated {len(products)} product documents in '{output_dir}' folder.")

if __name__ == "__main__":
    import sys
    
    # Check if a specific product name was provided
    if len(sys.argv) > 2 and sys.argv[1] == '--product':
        product_name = sys.argv[2]
        
        # Load products and find the specific product
        products = load_products()
        target_product = None
        
        for product in products:
            if product.get('Title', '').lower() == product_name.lower():
                target_product = product
                break
        
        if target_product:
            print(f"Regenerating document for: {product_name}")
            
            # Create output directory if it doesn't exist
            output_dir = "Product_Documents"
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # Generate document for the specific product
            filename = clean_filename(target_product['Title'])
            doc = create_product_document(target_product)
            output_path = os.path.join(output_dir, f"{filename}.docx")
            doc.save(output_path)
            
            print(f"Document regenerated: {output_path}")
        else:
            print(f"Product '{product_name}' not found in products database.")
            sys.exit(1)
    else:
        # Run the main function for all products
        main()
